from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    if level == 0:
        run.font.size = Pt(16)
    elif level == 1:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    return h


def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_list(doc: Document, items):
    for it in items or []:
        if not it:
            continue
        p = doc.add_paragraph(it)
        p.style = 'List Bullet'


def html_or_md_to_docx_bytes(content: str) -> bytes:
    # Very lightweight HTML -> DOCX mapper for the preview HTML we generate
    doc = Document()

    soup = BeautifulSoup(content or "", "lxml")
    body = soup.body or soup

    def walk(node):
        for el in node.children:
            if getattr(el, 'name', None) is None:
                # NavigableString
                continue
            name = el.name.lower()
            if name == 'h1':
                _add_heading(doc, el.get_text(strip=True), 0)
            elif name == 'h2':
                _add_heading(doc, el.get_text(strip=True), 1)
            elif name == 'h3':
                _add_heading(doc, el.get_text(strip=True), 2)
            elif name == 'p':
                _add_paragraph(doc, el.get_text(strip=True))
            elif name in ('ul', 'ol'):
                items = [li.get_text(strip=True) for li in el.find_all('li')]
                _add_list(doc, items)
            else:
                # recurse for sections/divs
                walk(el)

    walk(body)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
