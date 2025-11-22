from __future__ import annotations

import io
import re
import json
from typing import Optional, Tuple, Iterable

from bs4 import BeautifulSoup
from fastapi import UploadFile

from .schemas import Profile, Education, Experience, Course
from modules.shared.config import GPT_API_KEY, OPENAI_MODEL_PROFILE_EXTRACT, PROFILE_EXTRACT_LLM
from modules.shared.logging import get_logger

logger = get_logger("profile.extract")


def _read_bytes(file: UploadFile) -> bytes:
    return file.file.read() if hasattr(file, "file") else file.read()


def _ext(name: str | None) -> str:
    if not name:
        return ""
    i = name.rfind(".")
    return name[i + 1 :].lower() if i >= 0 else ""


def extract_text_from_upload(file: UploadFile) -> Tuple[str, list[str]]:
    """Return (text, notes). Supports pdf/docx/html; doc is best-effort.
    If unsupported, returns ("", [note]).
    """
    name = getattr(file, "filename", "") or "uploaded"
    ext = _ext(name)
    notes: list[str] = []
    data = _read_bytes(file) or b""
    if not data:
        return "", ["empty_file"]

    try:
        if ext == "pdf":
            try:
                from pdfminer.high_level import extract_text  # type: ignore

                text = extract_text(io.BytesIO(data)) or ""
                return text, notes
            except Exception as e:
                notes.append(f"pdfminer_error:{type(e).__name__}")
                # fallback very naive
                return "", notes
        elif ext == "docx":
            try:
                from docx import Document  # type: ignore

                doc = Document(io.BytesIO(data))
                parts: list[str] = []
                for p in doc.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)
                for table in getattr(doc, "tables", []):
                    for row in getattr(table, "rows", []):
                        for cell in getattr(row, "cells", []):
                            # A cell may contain nested paragraphs; keep simple newline separation
                            cell_text = "\n".join((para.text or "").strip() for para in cell.paragraphs if (para.text or "").strip())
                            if cell_text:
                                parts.append(cell_text)
                text = "\n".join(parts)
                return text, notes
            except Exception as e:
                notes.append(f"docx_error:{type(e).__name__}")
                return "", notes
        elif ext == "html" or ext == "htm":
            html = data.decode("utf-8", errors="ignore")
            soup = BeautifulSoup(html, "lxml")
            text = soup.get_text("\n", strip=True)
            return text, notes
        elif ext == "doc":
            # Best-effort: many environments cannot parse .doc reliably
            notes.append("doc_legacy_format_not_fully_supported")
            return "", notes
        else:
            notes.append("unsupported_extension")
            return "", notes
    except Exception as e:
        notes.append(f"extract_exception:{type(e).__name__}")
        return "", notes


SKILL_SEP = re.compile(r"[,/;]|\n|\t")


TECH_TOKENS = [
    "python", "java", "c++", "c#", "go", "typescript", "javascript", "react", "node", "spring",
    "docker", "kubernetes", "aws", "azure", "gcp", "sql", "postgres", "mysql", "mongodb",
    "terraform", "linux", "git", "ci/cd", "ros", "opencv", "pytorch", "tensorflow", "fastapi",
]


def _guess_skills(text: str) -> list[str]:
    # Prefer explicit Skills: section; otherwise only use known tokens
    parts: list[str] = []
    m = re.search(r"(?i)skills?\s*[:：]\s*(.+)$", text or "", flags=re.M)
    if m:
        line = m.group(1)
        parts.extend([s.strip() for s in re.split(SKILL_SEP, line) if s and len(s.strip()) > 1])
    # also scan full text for known tokens
    # also scan full text for known tokens
    low = (text or "").lower()
    for tok in TECH_TOKENS:
        if tok in low:
            parts.append(tok)
    # Normalize capitalization (simple heuristic)
    norm = []
    for s in parts:
        ss = s.strip()
        if not ss:
            continue
        if ss.isupper() and len(ss) <= 4:
            norm.append(ss)
        else:
            if ss in ["aws", "gcp", "sql", "ros", "ci/cd"]:
                norm.append(ss.upper())
            elif ss.lower() == "c++":
                norm.append("C++")
            elif ss.lower() == "c#":
                norm.append("C#")
            else:
                norm.append(ss.capitalize())
    # de-dup preserve order
    seen = set()
    out = []
    for s in norm:
        key = s.lower()
        if key in seen:
            continue
        seen.add(key)
        # filter out long Chinese sentences mistakenly captured
        if re.search(r"[\u4e00-\u9fff]", s or "") and len(s) > 12:
            continue
        out.append(s)
    return out[:64]


def _lines(text: str) -> Iterable[str]:
    for ln in (text or "").splitlines():
        s = ln.strip()
        if s:
            yield s


def parse_profile_rule_based(text: str) -> Profile:
    name = None
    contact = None
    location = None
    skills: list[str] = []
    education: list[Education] = []
    experience: list[Experience] = []
    projects: list[Experience] = []
    courses: list[Course] = []

    # Very light heuristics
    for i, ln in enumerate(_lines(text)):
        low = ln.lower()
        if i == 0 and len(ln.split()) <= 6 and not name:
            name = ln
            continue
        if ("email" in low or re.search(r"\b@\w+\.", ln)) and not contact:
            contact = ln
        if any(k in low for k in ["melbourne", "sydney", "beijing", "shanghai", "guangzhou", "shenzhen", "au", "nsw", "vic"]) and not location:
            location = ln
        if re.search(r"(?i)education|degree|bachelor|master|phd|university", ln):
            education.append(Education(school=ln))
        if re.search(r"(?i)^projects?\s*:\s*", ln):
            # split into bullets
            body = re.sub(r"(?i)^projects?\s*:\s*", "", ln)
            bs = [b.strip() for b in re.split(r"[•\-;]|\s\+\s|,\s(?=[A-Z])", body) if len(b.strip()) > 2]
            if bs:
                projects.append(Experience(company="Projects", bullets=bs[:12]))
            continue
        if re.search(r"(?i)^work\s+experience|^experience\s*:?$", low):
            # section header only
            continue
        if re.search(r"(?i)intern|engineer|developer", low):
            # try company - role pattern
            m = re.match(r"^(.{2,80}?)[\s\-|—|,|:|\|]+\s*([A-Za-z][^,•:|]{2,})$", ln)
            if m:
                experience.append(Experience(company=m.group(1).strip(), role=m.group(2).strip(), bullets=[]))
            else:
                # minimal line as company if short
                if len(ln) <= 80:
                    experience.append(Experience(company=ln, bullets=[]))

        # detect course codes in-line
        for m in _COURSE_CODE_ANY.finditer(ln.upper()):
            code = f"{m.group(1)}{m.group(2)}"
            # avoid duplicates
            if not any((c.code or "") == code for c in courses):
                courses.append(Course(code=code, name=None, topics=[], skills=[], tools=[]))

    # skills
    skills = _guess_skills(text)

    return Profile(
        name=name,
        contact=contact,
        location=location,
        education=education,
        experience=experience,
        projects=projects,
        courses=courses,
        skills=skills,
    )


def _shrink_text(text: str, max_len: int = 8000) -> str:
    if len(text) <= max_len:
        return text
    # Prefer lines containing key headers
    keys = re.compile(r"(?i)education|experience|project|skill|course|summary|profile|work|employment|intern|语言|教育|经历|项目|技能")
    lines = text.splitlines()
    picked: list[str] = []
    for ln in lines:
        if keys.search(ln or ""):
            picked.append(ln.strip())
        if len("\n".join(picked)) >= max_len:
            break
    if picked:
        s = "\n".join(picked)[:max_len]
        return s
    return text[:max_len]


def _llm_extract_profile(text: str) -> Optional[Profile]:
    if not GPT_API_KEY:
        return None
    if not text or len(text.strip()) < 10:
        return None
    try:
        from openai import OpenAI  # type: ignore

        client = OpenAI(api_key=GPT_API_KEY)
        system = (
            "Extract a structured resume Profile JSON from user-provided text. "
            "Do NOT fabricate facts. Only extract content that appears verbatim in the input. "
            "If unknown, leave fields empty. "
            "Output strictly JSON with keys: profile { name?, contact?, location?, education[], experience[], projects[], skills[], courses[], languages[], target_roles[], target_locations[] }."
        )
        shrunk = _shrink_text(text, max_len=8000)
        user = f"Text to extract from (may be Chinese/English):\n{shrunk}\nReturn only JSON, no markdown."

        resp = client.responses.create(
            model=OPENAI_MODEL_PROFILE_EXTRACT,
            input=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            timeout=15,
        )
        out = getattr(resp, "output_text", None) or getattr(resp, "content", "") or ""

        s = out.strip()
        if s.startswith("```"):
            s = "\n".join(s.splitlines()[1:])
            if s.endswith("```"):
                s = "\n".join(s.splitlines()[:-1])
        data = json.loads(s)
        # Normalize LLM output keys/types into our Profile schema
        prof_data_raw = data.get("profile", data)
        prof = _normalize_profile_dict(prof_data_raw)
        return prof
    except Exception as e:
        logger.warning(f"LLM extract failed: {e}")
        return None


def _as_list(v):
    if v is None:
        return []
    if isinstance(v, list):
        return v
    return [v]


_COURSE_CODE_ANY = re.compile(r"\b([A-Z]{3,4})\s?-?(\d{4})\b")


def _normalize_profile_dict(d: dict) -> Profile:
    name = d.get("name") or d.get("full_name")
    contact = d.get("contact") or d.get("contacts") or d.get("email")
    location = d.get("location") or d.get("city")

    # Education
    edus: list[Education] = []
    for item in _as_list(d.get("education")):
        if not item:
            continue
        if isinstance(item, str):
            edus.append(Education(school=item))
            continue
        if isinstance(item, dict):
            school = item.get("school") or item.get("institution") or item.get("university") or item.get("college") or item.get("name")
            degree = item.get("degree") or item.get("qualification")
            major = item.get("major") or item.get("field")
            start = item.get("start") or item.get("start_date")
            end = item.get("end") or item.get("end_date")
            if school:
                edus.append(Education(school=str(school), degree=degree, major=major, start=start, end=end))

    # Experience / Projects
    def _split_bullets(s: str) -> list[str]:
        parts = re.split(r"[\n•\-;]+|\s\+\s|,\s(?=[A-Z])", s)
        return [p.strip() for p in parts if len(p.strip()) > 2][:12]

    def norm_exps_and_projects() -> tuple[list[Experience], list[Experience]]:
        out_exp: list[Experience] = []
        out_proj: list[Experience] = []
        # normalize dict items
        for key in ("experience", "projects"):
            for item in _as_list(d.get(key)):
                if not item:
                    continue
                if isinstance(item, dict):
                    company = item.get("company") or item.get("employer") or item.get("org") or item.get("organization") or item.get("name")
                    role = item.get("role") or item.get("position") or item.get("title")
                    bullets = _as_list(item.get("bullets") or item.get("highlights") or item.get("responsibilities") or item.get("achievements"))
                    bullets = [str(b).strip() for b in bullets if str(b).strip()]
                    if company or role or bullets:
                        (out_proj if key == "projects" else out_exp).append(Experience(company=company or "", role=role, bullets=bullets))
                elif isinstance(item, str):
                    s = item.strip()
                    s_low = s.lower()
                    if s_low.startswith("project:") or s_low.startswith("projects:"):
                        out_proj.append(Experience(company="Projects", role=None, bullets=_split_bullets(s.split(":",1)[1])))
                    elif s_low.startswith("work experience") or s_low == "experience":
                        # section headers: ignore
                        continue
                    else:
                        # pattern: Company - Role
                        m = re.match(r"^([^\-|—|,|:|\|]{2,})\s*[\-|—|,|:|\|]+\s*([^,•|:]{2,})$", s)
                        if m:
                            out_exp.append(Experience(company=m.group(1).strip(), role=m.group(2).strip(), bullets=[]))
                        else:
                            # if contains many tech tokens, treat as project bullets
                            if any(tok in s_low for tok in TECH_TOKENS):
                                out_proj.append(Experience(company="Projects", role=None, bullets=_split_bullets(s)))
                            else:
                                # fallback minimal exp line as company
                                if len(s) <= 80 and not re.search(r"[\u4e00-\u9fff]{8,}", s):
                                    out_exp.append(Experience(company=s))
        return out_exp, out_proj

    exps, projs = norm_exps_and_projects()

    # Skills
    skills_in = _as_list(d.get("skills"))
    skills_str = []
    for v in skills_in:
        if isinstance(v, str):
            skills_str.append(v)
        elif isinstance(v, dict):
            n = v.get("name")
            if n:
                skills_str.append(str(n))
    skills = _guess_skills("\n".join(skills_str)) if skills_str else []

    # Courses
    courses: list[Course] = []
    for item in _as_list(d.get("courses")):
        if not item:
            continue
        if isinstance(item, str):
            m = _COURSE_CODE_ANY.search(item.upper())
            if m:
                code = f"{m.group(1)}{m.group(2)}"
                courses.append(Course(code=code, name=None, topics=[], skills=[], tools=[]))
            else:
                courses.append(Course(code=None, name=item, topics=[], skills=[], tools=[]))
        elif isinstance(item, dict):
            code = item.get("code")
            name_c = item.get("name")
            if not code and isinstance(name_c, str):
                m = _COURSE_CODE_ANY.search(name_c.upper())
                if m:
                    code = f"{m.group(1)}{m.group(2)}"
            topics = _as_list(item.get("topics"))
            skills_c = _as_list(item.get("skills"))
            tools = _as_list(item.get("tools"))
            courses.append(Course(code=code, name=name_c, topics=[str(x) for x in topics if str(x).strip()], skills=[str(x) for x in skills_c if str(x).strip()], tools=[str(x) for x in tools if str(x).strip()]))

    # Languages / targets
    languages = [str(x) for x in _as_list(d.get("languages")) if str(x).strip()]
    target_roles = [str(x) for x in _as_list(d.get("target_roles")) if str(x).strip()]
    target_locations = [str(x) for x in _as_list(d.get("target_locations")) if str(x).strip()]

    return Profile(
        name=name,
        contact=contact,
        location=location,
        summary=d.get("summary") or d.get("profile_summary"),
        education=edus,
        experience=exps,
        projects=projs,
        skills=skills,
        courses=courses,
        languages=languages,
        target_roles=target_roles,
        target_locations=target_locations,
    )


def merge_profiles(primary: Profile, secondary: Profile) -> Profile:
    """Merge secondary into primary; keep primary values if present. Lists merged unique."""
    def pick(a, b):
        return a if (a and str(a).strip()) else b

    def merge_list(x: list[str], y: list[str]) -> list[str]:
        seen = set()
        out = []
        for v in (x or []) + (y or []):
            if not v:
                continue
            k = v.strip().lower()
            if k in seen:
                continue
            seen.add(k)
            out.append(v)
        return out

    p = Profile(
        name=pick(primary.name, secondary.name),
        contact=pick(primary.contact, secondary.contact),
        location=pick(primary.location, secondary.location),
        summary=pick(primary.summary, secondary.summary),
        education=primary.education or secondary.education,
        experience=primary.experience or secondary.experience,
        projects=primary.projects or secondary.projects,
        skills=merge_list(primary.skills or [], secondary.skills or []),
        courses=primary.courses or secondary.courses,
        languages=merge_list(primary.languages or [], secondary.languages or []),
        target_roles=merge_list(primary.target_roles or [], secondary.target_roles or []),
        target_locations=merge_list(primary.target_locations or [], secondary.target_locations or []),
    )
    return p


def analyze_text_to_profile(text: str) -> Tuple[Profile, list[str]]:
    notes: list[str] = []
    prof = None
    if GPT_API_KEY and PROFILE_EXTRACT_LLM != "0":
        prof = _llm_extract_profile(text)
    if prof:
        notes.append("llm_ok")
        return prof, notes
    # fallback
    notes.append("rule_based")
    return parse_profile_rule_based(text), notes


def englishize_skills(skills: list[str]) -> list[str]:
    # Preserve existing phrases; title-case simple tokens; ensure acronyms
    out = []
    for s in skills or []:
        if not s:
            continue
        t = s.strip()
        low = t.lower()
        if low in {"sql", "nlp"}:
            out.append(low.upper())
            continue
        if low in {"c++", "c#"}:
            out.append(low.upper())
            continue
        # If phrase contains separators or spaces, assume well-formed and keep
        if any(ch in t for ch in [" ", "&", "+", "/"]):
            out.append(t)
        else:
            # title-case single tokens
            out.append(t.title())
    # de-dup case-insensitive
    seen = set()
    uniq = []
    for s in out:
        k = s.lower()
        if k in seen:
            continue
        seen.add(k)
        uniq.append(s)
    return uniq
